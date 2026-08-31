#!/usr/bin/env python3
import json,re,hashlib,subprocess,sys,zipfile,xml.etree.ElementTree as ET
from pathlib import Path
from docx import Document
ROOT=Path(__file__).resolve().parents[1];SRC=ROOT/'.audit-source'/'kk';OUT=ROOT/'zaytoona'/'kefayat'
SUBJECTS={'arabic':'اللغة العربية','math':'الرياضيات','islamic':'التربية الإسلامية','nurturing':'التنشئة/العلوم الحياتية'};EXPECTED={(g,s) for g in range(1,5) for s in SUBJECTS}
def norm(x):return re.sub(r'\s+',' ',str(x or '').replace('\xa0',' ')).strip()
def sha(x):return hashlib.sha256(x if isinstance(x,(bytes,bytearray)) else x.encode()).hexdigest()
def grade(name):
 n=name.replace('أ','ا').replace('إ','ا').replace('ى','ي')
 for w,g in [('الرابع',4),('رابع',4),('الثالث',3),('ثالث',3),('الثاني',2),('ثاني',2),('الأول',1),('الاول',1),('اول',1),('أول',1)]:
  if w in n:return g
 m=re.search(r'(?:الصف|صف)\s*([1-4١٢٣٤])|الصف([1-4١٢٣٤])|صف([1-4١٢٣٤])',n) or re.search(r'([1-4١٢٣٤])',n)
 if not m:return None
 x=next(v for v in m.groups() if v);return {'١':1,'٢':2,'٣':3,'٤':4}.get(x,int(x) if x.isdigit() else None)
def subject(name):
 if 'العربية' in name:return 'arabic'
 if 'الرياضيات' in name:return 'math'
 if 'اسلامية' in name or 'إسلامية' in name:return 'islamic'
 if 'تنشئة' in name or 'علوم' in name or 'وطنية' in name:return 'nurturing'
 return None
def hscore(c):return sum(any(k in x for k in ['كفاية','الكفاية','معيار','المعيار','مؤشر','المؤشر','قيمة','القيم','مجال','المجال','نواتج','نتاج']) for x in c)
def fields(h,c):
 r={}
 for a,b in zip(h,c):
  a=norm(a);b=norm(b)
  if not a:continue
  if 'مجال' in a:r['domain']=b
  elif 'كفاية' in a:r['competency']=b
  elif 'معيار' in a:r['standard']=b
  elif 'مؤشر' in a or 'نتاج' in a:r['indicator']=b
  elif 'قيمة' in a:r['value']=b
  else:r.setdefault('extra',[]).append({'header':a,'text':b})
 return r
def xml_text(path):
 try:
  with zipfile.ZipFile(path) as z:
   root=ET.fromstring(z.read('word/document.xml'));ns={'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'};out=[]
   for p in root.findall('.//w:p',ns):
    t=norm(''.join(x.text or '' for x in p.findall('.//w:t',ns)))
    if t:out.append(t)
   return out
 except Exception:return []
def parse_doc(path):
 g=grade(path.name);s=subject(path.name);doc=Document(path);records=[];xt=xml_text(path);diag={'file':path.name,'grade':g,'subject':s,'tables':[],'paragraphs':[norm(p.text) for p in doc.paragraphs if norm(p.text)][:20],'xmlText':xt[:250]}
 if not g or not s:return [],{**diag,'status':'UNMAPPED'}
 for ti,t in enumerate(doc.tables):
  rows=[[norm(c.text) for c in row.cells] for row in t.rows];rows=[r for r in rows if any(r)];diag['tables'].append({'table':ti,'rowCount':len(rows),'colCount':max([len(r) for r in rows] or [0]),'preview':rows[:8]})
  if not rows:continue
  hi=max(range(min(len(rows),8)),key=lambda i:hscore(rows[i]));h=rows[hi]
  if hscore(h)<1:h=[f'column_{i+1}' for i in range(len(h))];hi=-1
  for ri,c in enumerate(rows[hi+1:],hi+1):
   if not any(c) or len(c)!=len(h):continue
   f=fields(h,c)
   if not any(f.get(k) for k in ['competency','standard','indicator','value','domain']):continue
   raw=' | '.join(c);key=sha(f'{s}|{g}|{raw}')
   records.append({'id':'KK-'+key[:16],'grade':g,'subject':s,'subjectLabel':SUBJECTS[s],'domain':f.get('domain',''),'competency':f.get('competency',''),'standard':f.get('standard',''),'indicator':f.get('indicator',''),'value':f.get('value',''),'extra':f.get('extra',[]),'raw':raw,'source':{'repository':'smileeyes1/kk','branch':'main','file':path.name,'table':ti,'row':ri,'sourceSha':sha(path.read_bytes())},'provenance':'USER_PROVIDED_REFERENCE','officialStatus':'UNVERIFIED'})
 if not records and xt:
  for i,line in enumerate(xt):
   if any(k in line for k in ['الكفاية','كفاية']) and not any(k in line for k in ['الرئيسية','الرئيسة','معايير','معاييرها']):
    key=sha(f'{s}|{g}|xml|{i}|{line}');records.append({'id':'KK-'+key[:16],'grade':g,'subject':s,'subjectLabel':SUBJECTS[s],'domain':'','competency':line,'standard':'','indicator':'','value':'','extra':[],'raw':line,'source':{'repository':'smileeyes1/kk','branch':'main','file':path.name,'xmlParagraph':i,'sourceSha':sha(path.read_bytes())},'provenance':'USER_PROVIDED_REFERENCE','officialStatus':'UNVERIFIED'})
 return records,{**diag,'status':'OK','records':len(records)}
def main():
 if not SRC.exists():SRC.parent.mkdir(parents=True,exist_ok=True);subprocess.run(['git','clone','--depth','1','https://github.com/smileeyes1/kk.git',str(SRC)],check=True)
 files=list(SRC.glob('*.docx'));allr=[];man=[]
 for p in files:r,m=parse_doc(p);allr+=r;man.append(m)
 cov={(g,s):0 for g,s in EXPECTED}
 for r in allr:cov[(r['grade'],r['subject'])]+=1
 seen={};dups=[]
 for r in allr:
  k=(r['grade'],r['subject'],r['domain'],r['competency'],r['standard'],r['indicator'],r['value'])
  if k in seen:dups.append({'first':seen[k],'duplicate':r['id'],'key':k})
  else:seen[k]=r['id']
 gaps=[{'grade':g,'subject':s,'count':cov[(g,s)]} for g,s in sorted(EXPECTED) if cov[(g,s)]==0]
 variants={};conf=[]
 for r in allr:
  k=(r['grade'],r['subject'],r['competency'])
  if r['competency']:variants.setdefault(k,set()).add((r['standard'],r['indicator'],r['value']))
 for k,v in variants.items():
  if len(v)>1:conf.append({'key':k,'variants':list(v)[:20]})
 unique={}
 for r in allr:unique.setdefault((r['grade'],r['subject'],r['domain'],r['competency'],r['standard'],r['indicator'],r['value']),r)
 records=list(unique.values());OUT.mkdir(parents=True,exist_ok=True)
 cat={'schemaVersion':5,'audit':'Ω KEFAYAT MASTER AUDIT & COMPLETION','generatedFrom':{'repository':'smileeyes1/kk','branch':'main'},'officialityPolicy':'USER-PROVIDED REFERENCE; NOT OFFICIAL-VERIFIED','recordCount':len(records),'inputRecordCount':len(allr),'deduplicated':len(allr)-len(records),'coverage':{'grades':[1,2,3,4],'subjects':list(SUBJECTS),'matrix':{f'{g}|{s}':cov[(g,s)] for g,s in sorted(EXPECTED)}},'records':records}
 (OUT/'master-audit.json').write_text(json.dumps({'status':'NO-GO' if gaps else 'PASS','files':man,'coverage':cat['coverage'],'duplicates':dups,'conflicts':conf,'gaps':gaps,'inputRecordCount':len(allr),'canonicalRecordCount':len(records)},ensure_ascii=False,indent=2)+'\n',encoding='utf8')
 (OUT/'catalog.json').write_text(json.dumps(cat,ensure_ascii=False,indent=2)+'\n',encoding='utf8');(OUT/'catalog.min.json').write_text(json.dumps(cat,ensure_ascii=False,separators=(',',':'))+'\n',encoding='utf8')
 out={'status':'PASS' if not gaps else 'NO-GO','files':len(files),'records':len(records),'duplicates':len(dups),'conflicts':len(conf),'gaps':gaps,'coverage':cat['coverage']['matrix']};print(json.dumps(out,ensure_ascii=False,indent=2));
 if gaps:sys.exit(2)
if __name__=='__main__':main()
